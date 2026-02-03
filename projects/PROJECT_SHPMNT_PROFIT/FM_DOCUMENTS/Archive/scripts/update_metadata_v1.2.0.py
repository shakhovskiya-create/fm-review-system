#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт обновления метаданных документа FM-LS-PROFIT v1.2.0
"""

from docx import Document
from datetime import datetime
import re

def update_version_metadata(doc, old_version, new_version):
    """Обновить версию во всех местах документа"""
    updates = []

    # 1. Header/Footer
    print("1. Обновление header/footer...")
    for section in doc.sections:
        # Header
        for para in section.header.paragraphs:
            if old_version in para.text:
                old_text = para.text
                para.text = para.text.replace(old_version, new_version)
                updates.append(f"Header: {old_text[:50]} → {para.text[:50]}")
                print(f"   ✓ Header обновлен")

        # Footer
        for para in section.footer.paragraphs:
            if old_version in para.text:
                old_text = para.text
                para.text = para.text.replace(old_version, new_version)
                updates.append(f"Footer: {old_text[:50]} → {para.text[:50]}")
                print(f"   ✓ Footer обновлен")

    # 2. Первые 50 параграфов (титульная страница, заголовки)
    print("\n2. Обновление титульной страницы...")
    for i in range(min(50, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if old_version in para.text:
            old_text = para.text
            para.text = para.text.replace(old_version, new_version)
            updates.append(f"Para {i}: {old_text[:50]} → {para.text[:50]}")
            print(f"   ✓ Параграф {i} обновлен")

    # 3. Поиск по всему документу
    print("\n3. Поиск и замена по всему документу...")
    for i, para in enumerate(doc.paragraphs):
        if old_version in para.text:
            old_text = para.text
            # Заменяем все вхождения
            new_text = para.text.replace(old_version, new_version)
            para.clear()
            para.add_run(new_text)
            if i >= 50:  # Уже обновили первые 50
                updates.append(f"Para {i}: версия обновлена")
                print(f"   ✓ Параграф {i} обновлен")

    # 4. Обновление даты
    print("\n4. Обновление даты последнего изменения...")
    today = datetime.now().strftime("%d.%m.%Y")
    found_date = False

    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.lower()
        if ('дата' in text or 'date' in text) and ('изменени' in text or 'версии' in text or 'последн' in text):
            # Ищем дату в формате DD.MM.YYYY
            date_pattern = r'\d{2}\.\d{2}\.\d{4}'
            if re.search(date_pattern, para.text):
                old_text = para.text
                new_text = re.sub(date_pattern, today, para.text)
                para.clear()
                para.add_run(new_text)
                updates.append(f"Дата: {old_text[:50]} → {new_text[:50]}")
                print(f"   ✓ Дата обновлена в параграфе {i}")
                found_date = True
                break

    if not found_date:
        # Добавим дату в начало документа
        print("   ⚠️  Дата не найдена, добавляем новый параграф")
        # Ищем место после заголовка
        for i in range(10):
            if doc.paragraphs[i].text.strip():
                new_para = doc.paragraphs[i]._element
                added = doc.add_paragraph()._element
                new_para.addnext(added)
                doc.paragraphs[i+1].add_run(f"Дата последнего изменения: {today}")
                updates.append(f"Дата добавлена после параграфа {i}")
                print(f"   ✓ Дата добавлена после параграфа {i}")
                break

    return updates

def main():
    print("=" * 70)
    print("ОБНОВЛЕНИЕ МЕТАДАННЫХ FM-LS-PROFIT v1.2.0")
    print("=" * 70)
    print()

    input_file = '/Users/antonsahovskii/Documents/claude-agents/fm-review-system/PROJECT_SHPMNT_PROFIT/FM_DOCUMENTS/FM-LS-PROFIT-v1.2.0.docx'

    print(f"📖 Загрузка документа: {input_file}")
    doc = Document(input_file)
    print(f"   Всего параграфов: {len(doc.paragraphs)}")
    print(f"   Всего секций: {len(doc.sections)}")
    print()

    # Обновляем версию
    updates = update_version_metadata(doc, 'v1.1.0', 'v1.2.0')

    # Также пробуем без 'v'
    print("\n5. Дополнительная проверка версий без 'v'...")
    for i, para in enumerate(doc.paragraphs[:100]):
        if '1.1.0' in para.text and 'v1.1.0' not in para.text:
            old_text = para.text
            para.text = para.text.replace('1.1.0', '1.2.0')
            updates.append(f"Para {i}: 1.1.0 → 1.2.0")
            print(f"   ✓ Параграф {i} обновлен (без 'v')")

    # Добавляем запись в историю изменений (если есть такой раздел)
    print("\n6. Поиск раздела 'История изменений'...")
    history_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        if 'истори' in text and 'изменени' in text:
            history_idx = i
            print(f"   ✓ Раздел найден на параграфе {i}")
            break

    if history_idx:
        # Добавляем запись о версии 1.2.0
        today = datetime.now().strftime("%d.%m.%Y")
        new_para = doc.paragraphs[history_idx]._element
        added = doc.add_paragraph()._element
        new_para.addnext(added)

        new_para_obj = doc.paragraphs[history_idx + 1]
        new_para_obj.add_run(f"\nВерсия 1.2.0 ({today}): ").bold = True
        new_para_obj.add_run("Интеграция 22 логических исправлений из аудита Logic Review. Критические исправления: защита от race condition, уточнение формул, правила округления, фиксация потребности, пересчет санкций при возвратах, автопродление, разделение фиксации. Высокоприоритетные: аннулирование при изменении ЛС, атомарная проверка, push-уведомления, реабилитация через 3 месяца, эскалация НПСС, рентабельность возврата, мульти-БЮ, лимиты, разделение резервов. Средние: автоснятие блокировки, исключение сторнированных РТУ, процентный контроль убытка, РТУ без повтора. Низкие: границы KPI, FAQ автопродление.")

        updates.append(f"История изменений: добавлена запись v1.2.0")
        print(f"   ✓ Добавлена запись в историю изменений")
    else:
        print("   ⚠️  Раздел 'История изменений' не найден")

    # Сохранение
    print("\n" + "=" * 70)
    print("СОХРАНЕНИЕ РЕЗУЛЬТАТА")
    print("=" * 70)
    print(f"\n📝 Сохранение документа: {input_file}")
    doc.save(input_file)
    print("✓ Документ успешно сохранен")

    # Отчет
    print("\n" + "=" * 70)
    print("ИТОГОВЫЙ ОТЧЕТ")
    print("=" * 70)
    print(f"\n✅ Всего обновлений: {len(updates)}")
    for update in updates[:10]:  # Показываем первые 10
        print(f"   • {update}")
    if len(updates) > 10:
        print(f"   ... и еще {len(updates) - 10} обновлений")

    print("\n✅ Метаданные документа v1.2.0 обновлены!")
    print("=" * 70)

if __name__ == "__main__":
    main()
