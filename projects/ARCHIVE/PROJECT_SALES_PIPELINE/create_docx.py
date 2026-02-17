#!/usr/bin/env python3
"""
Скрипт для создания профессиональных DOCX версий мини-ФМ из Markdown файлов
Версия 2.0 — улучшенное форматирование
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Путь к папке
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_DIR = os.path.join(BASE_DIR, 'docx')

# Файлы для конвертации
FM_FILES = [
    'FM-001-PIPELINE-SPLIT.md',
    'FM-002-STAGES.md',
    'FM-003-ANALOG-REQUEST.md',
    'FM-004-AUTO-CLOSE.md',
    'FM-005-MASTER-ESTIMATE.md',
]

# Цвета
CORP_BLUE = RGBColor(0, 82, 147)  # Корпоративный синий
CORP_LIGHT_BLUE = RGBColor(217, 226, 243)  # Светло-синий для заливки
CORP_GRAY = RGBColor(89, 89, 89)  # Серый для текста
HEADER_BG = RGBColor(0, 82, 147)  # Синий фон для заголовков таблиц


def set_cell_shading(cell, color):
    """Установить заливку ячейки"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, border_color="000000", border_size=4):
    """Установить границы ячейки"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_line(doc):
    """Добавить горизонтальную линию"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Создаем линию через border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '0052A3')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_markdown(content):
    """Парсинг markdown в структуру"""
    lines = content.split('\n')
    elements = []
    in_code_block = False
    code_block_content = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Код блоки
        if line.strip().startswith('```'):
            if in_code_block:
                elements.append({'type': 'code', 'content': '\n'.join(code_block_content)})
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Заголовки
        if line.startswith('# '):
            elements.append({'type': 'h1', 'content': line[2:].strip()})
        elif line.startswith('## '):
            elements.append({'type': 'h2', 'content': line[3:].strip()})
        elif line.startswith('### '):
            elements.append({'type': 'h3', 'content': line[4:].strip()})
        elif line.startswith('#### '):
            elements.append({'type': 'h4', 'content': line[5:].strip()})

        # Таблицы
        elif line.strip().startswith('|') and '|' in line[1:]:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'content': table_lines})
            continue

        # Горизонтальная линия
        elif line.strip() == '---':
            elements.append({'type': 'hr', 'content': ''})

        # Блок цитаты (информационный блок)
        elif line.strip().startswith('>'):
            quote_lines = [line.strip()[1:].strip()]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            elements.append({'type': 'quote', 'content': '\n'.join(quote_lines)})
            continue

        # Списки
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append({'type': 'list_item', 'content': line.strip()[2:]})
        elif re.match(r'^\d+\.\s', line.strip()):
            elements.append({'type': 'numbered_item', 'content': re.sub(r'^\d+\.\s', '', line.strip())})

        # Чек-боксы
        elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
            checked = '[x]' in line
            text = line.strip()[6:]
            elements.append({'type': 'checkbox', 'content': text, 'checked': checked})

        # Обычный текст
        elif line.strip():
            elements.append({'type': 'paragraph', 'content': line.strip()})

        i += 1

    return elements


def format_text_with_styles(paragraph, text):
    """Форматирование текста с поддержкой bold и code"""
    # Разбиваем текст по маркерам форматирования
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Arial'
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = CORP_BLUE
        elif part:
            run = paragraph.add_run(part)
            run.font.name = 'Arial'


def create_docx(elements, output_path, title, fm_code):
    """Создание профессионального DOCX из элементов"""
    doc = Document()

    # === НАСТРОЙКА ДОКУМЕНТА ===

    # Размеры страницы и поля
    section = doc.sections[0]
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # === СТИЛИ ===

    # Стиль Normal
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = CORP_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Стиль для заголовка документа
    try:
        title_style = doc.styles.add_style('DocTitle', WD_STYLE_TYPE.PARAGRAPH)
    except:
        title_style = doc.styles['DocTitle']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = CORP_BLUE
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === КОЛОНТИТУЛЫ ===

    # Верхний колонтитул
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run(f'{fm_code}')
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.color.rgb = CORP_BLUE
    run.bold = True

    # Нижний колонтитул
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run1 = footer_para.add_run('Версия 1.0.0 | ')
    run1.font.name = 'Arial'
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor(128, 128, 128)

    run2 = footer_para.add_run(datetime.now().strftime('%d.%m.%Y'))
    run2.font.name = 'Arial'
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    run3 = footer_para.add_run(' | Статус: ПРЕДЛОЖЕНИЕ')
    run3.font.name = 'Arial'
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor(128, 128, 128)

    # === СОДЕРЖИМОЕ ===

    section_number = 0

    for elem in elements:
        if elem['type'] == 'h1':
            # Главный заголовок
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(elem['content'])
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = CORP_BLUE

            # Линия под заголовком
            add_horizontal_line(doc)

        elif elem['type'] == 'h2':
            section_number += 1
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)

            # Номер секции
            run_num = p.add_run(f'{section_number}. ')
            run_num.font.name = 'Arial'
            run_num.font.size = Pt(14)
            run_num.font.bold = True
            run_num.font.color.rgb = CORP_BLUE

            # Текст заголовка
            run = p.add_run(elem['content'])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = CORP_BLUE

        elif elem['type'] == 'h3':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(elem['content'])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = CORP_GRAY

        elif elem['type'] == 'h4':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(elem['content'])
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = CORP_GRAY

        elif elem['type'] == 'paragraph':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            format_text_with_styles(p, elem['content'])

        elif elem['type'] == 'quote':
            # Информационный блок с заливкой
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)

            # Добавляем заливку через XML
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D9E2F3')  # Светло-синий
            pPr.append(shd)

            run = p.add_run(elem['content'])
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = CORP_GRAY

        elif elem['type'] == 'list_item':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(3)

            run_bullet = p.add_run('• ')
            run_bullet.font.name = 'Arial'
            run_bullet.font.color.rgb = CORP_BLUE
            run_bullet.bold = True

            format_text_with_styles(p, elem['content'])

        elif elem['type'] == 'numbered_item':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(3)
            format_text_with_styles(p, elem['content'])

        elif elem['type'] == 'checkbox':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(3)

            prefix = '☑' if elem.get('checked') else '☐'
            run_check = p.add_run(f'{prefix} ')
            run_check.font.name = 'Arial'
            run_check.font.size = Pt(12)
            if elem.get('checked'):
                run_check.font.color.rgb = RGBColor(0, 128, 0)  # Зеленый
            else:
                run_check.font.color.rgb = CORP_GRAY

            format_text_with_styles(p, elem['content'])

        elif elem['type'] == 'code':
            # Код блок с рамкой
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)

            # Заливка для кода
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F5F5F5')  # Светло-серый
            pPr.append(shd)

            run = p.add_run(elem['content'])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = CORP_GRAY

        elif elem['type'] == 'table':
            # Парсим таблицу
            rows = []
            for line in elem['content']:
                if '---' in line and line.replace('-', '').replace('|', '').strip() == '':
                    continue
                cells = [c.strip() for c in line.strip('|').split('|')]
                if cells and any(c for c in cells):
                    rows.append(cells)

            if rows:
                # Определяем максимальное количество колонок
                max_cols = max(len(row) for row in rows)

                # Выравниваем все строки
                for row in rows:
                    while len(row) < max_cols:
                        row.append('')

                table = doc.add_table(rows=len(rows), cols=max_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Ширина таблицы
                table.autofit = False
                table_width = Cm(16)  # Общая ширина
                col_width = table_width / max_cols

                for i, row_data in enumerate(rows):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            cell = table.rows[i].cells[j]
                            cell.width = col_width
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                            # Очищаем и добавляем текст
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(4)
                            p.paragraph_format.space_after = Pt(4)

                            format_text_with_styles(p, cell_data)

                            # Границы
                            set_cell_borders(cell, "0052A3", 6)

                            # Первая строка - заголовок
                            if i == 0:
                                set_cell_shading(cell, (0, 82, 147))  # Синий фон
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # Белый текст
                            else:
                                # Чередование строк
                                if i % 2 == 0:
                                    set_cell_shading(cell, (245, 245, 245))  # Светло-серый

                # Отступ после таблицы
                doc.add_paragraph()

        elif elem['type'] == 'hr':
            add_horizontal_line(doc)

    # === СОХРАНЕНИЕ ===
    doc.save(output_path)
    print(f"✓ Создан: {output_path}")


def main():
    os.makedirs(DOCX_DIR, exist_ok=True)

    print("\n" + "="*60)
    print("  ГЕНЕРАЦИЯ DOCX ФАЙЛОВ")
    print("="*60 + "\n")

    for md_file in FM_FILES:
        md_path = os.path.join(BASE_DIR, md_file)
        if not os.path.exists(md_path):
            print(f"✗ Файл не найден: {md_path}")
            continue

        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        elements = parse_markdown(content)

        # Извлекаем заголовок и код ФМ
        title = md_file.replace('.md', '')
        fm_code = md_file.replace('.md', '').upper()

        for elem in elements:
            if elem['type'] == 'h1':
                title = elem['content']
                break

        docx_file = md_file.replace('.md', '.docx')
        docx_path = os.path.join(DOCX_DIR, docx_file)

        create_docx(elements, docx_path, title, fm_code)

    print("\n" + "="*60)
    print(f"  Готово! Файлы в папке: {DOCX_DIR}")
    print("="*60 + "\n")


if __name__ == '__main__':
    main()
