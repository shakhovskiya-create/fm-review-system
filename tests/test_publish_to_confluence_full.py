"""
Comprehensive unit tests for scripts/publish_to_confluence.py

Target: 95-100% coverage. Mocks external deps (docx, ConfluenceClient, network).
"""
import os
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch, mock_open

import pytest

# Add scripts to path (conftest does this too)
SCRIPTS_DIR = Path(__file__).parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))


# ── get_cell_color (requires qn from docx; mock when docx unavailable) ─────

class TestGetCellColor:
    """Test get_cell_color - mock qn/tc structure when docx may be unavailable."""

    def test_returns_color_when_shd_found(self):
        from publish_to_confluence import get_cell_color, qn
        if qn is None:
            pytest.skip("docx not installed - qn is None")
        cell = MagicMock()
        shd = MagicMock()
        shd.get.return_value = "FFDD00"
        tcPr = MagicMock()
        tcPr.find.return_value = shd
        cell._tc.find.return_value = tcPr
        result = get_cell_color(cell)
        assert result == "FFDD00"

    def test_returns_none_when_tcPr_is_none(self):
        from publish_to_confluence import get_cell_color, qn
        if qn is None:
            pytest.skip("docx not installed")
        cell = MagicMock()
        cell._tc.find.return_value = None
        assert get_cell_color(cell) is None

    def test_returns_none_when_shd_is_none(self):
        from publish_to_confluence import get_cell_color, qn
        if qn is None:
            pytest.skip("docx not installed")
        cell = MagicMock()
        tcPr = MagicMock()
        tcPr.find.return_value = None
        cell._tc.find.return_value = tcPr
        assert get_cell_color(cell) is None

    def test_is_warning_table_works_with_patched_get_cell_color(self):
        """When docx unavailable, is_warning_table uses get_cell_color - patch it."""
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Normal 1x1 text"
            assert is_warning_table(table) is None


# ── _get_page_id from project file ─────────────────────────────────────────

class TestGetPageIdFromFile:
    """Test _get_page_id reading from projects/PROJECT/CONFLUENCE_PAGE_ID."""

    def test_reads_from_project_file_when_exists(self, tmp_path):
        from publish_to_confluence import _get_page_id
        project_dir = tmp_path / "projects" / "MY_PROJECT"
        project_dir.mkdir(parents=True)
        (project_dir / "CONFLUENCE_PAGE_ID").write_text("12345678\n")
        script_path = tmp_path / "scripts" / "publish_to_confluence.py"
        script_path.parent.mkdir(parents=True, exist_ok=True)
        with patch.dict(os.environ, {"CONFLUENCE_PAGE_ID": ""}, clear=False):
            with patch("publish_to_confluence.__file__", str(script_path)):
                with patch("publish_to_confluence.os.path.dirname") as mock_dirname:
                    with patch("publish_to_confluence.os.path.abspath", return_value=str(script_path)):
                        def dirname_side_effect(p):
                            if p == str(script_path):
                                return str(tmp_path / "scripts")
                            if "scripts" in str(p):
                                return str(tmp_path)
                            return os.path.dirname(p)
                        mock_dirname.side_effect = dirname_side_effect
                        result = _get_page_id("MY_PROJECT")
        assert result == "12345678"

    def test_skips_comment_lines_in_page_id_file(self):
        from publish_to_confluence import _get_page_id
        with patch("publish_to_confluence.os.path.dirname") as mock_dn:
            with patch("publish_to_confluence.os.path.abspath"):
                with patch("publish_to_confluence.os.path.isfile", return_value=True):
                    with patch("builtins.open", mock_open(read_data="# comment\n\n55555555\n")):
                        mock_dn.return_value = "/tmp"
                        result = _get_page_id("PROJ")
        assert result == "55555555"

    def test_skips_non_digit_lines(self):
        from publish_to_confluence import _get_page_id
        with patch("publish_to_confluence.os.path.dirname") as mock_dn:
            with patch("publish_to_confluence.os.path.abspath"):
                with patch("publish_to_confluence.os.path.isfile", return_value=True):
                    with patch("builtins.open", mock_open(read_data="abc\n99999999\n")):
                        mock_dn.return_value = "/tmp"
                        result = _get_page_id("X")
        assert result == "99999999"


# ── is_warning_table (with patched get_cell_color) ─────────────────────────

class TestIsWarningTableFull:
    """Test is_warning_table branches - patch get_cell_color to avoid qn."""

    def test_critical_keyword_returns_warning(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Критичная зависимость от внешней системы"
            assert is_warning_table(table) == "warning"

    def test_dependency_keyword_returns_warning(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Critical dependency"
            assert is_warning_table(table) == "warning"

    def test_note_keyword_returns_note(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Исключение: при отключенном модуле"
            assert is_warning_table(table) == "note"

    def test_attention_keyword_returns_note(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Важно: проверить настройки"
            assert is_warning_table(table) == "note"

    def test_emoji_keyword_returns_note(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            row = MagicMock()
            cell = MagicMock()
            cell.text = "Внимание: требуется проверка"
            row.cells = [cell]
            table.rows = [row]
            assert is_warning_table(table) == "note"

    def test_normal_1x1_returns_none(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import is_warning_table
            table = MagicMock()
            table.rows = [MagicMock()]
            table.rows[0].cells = [MagicMock()]
            table.rows[0].cells[0].text = "Just a normal cell"
            assert is_warning_table(table) is None


# ── table_to_html (emoji stripping, cell colors) ────────────────────────────

class TestTableToHtmlFull:
    """Test table_to_html emoji stripping and cell color handling."""

    def _make_table(self, rows_data):
        rows = []
        for row_cells in rows_data:
            row = MagicMock()
            cells = []
            for text in row_cells:
                cell = MagicMock()
                cell.text = text
                cell._tc = MagicMock()
                cell._tc.find.return_value = None
                cells.append(cell)
            row.cells = cells
            rows.append(row)
        table = MagicMock()
        table.rows = rows
        return table

    def test_warning_panel_strips_emoji(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import table_to_html
            table = self._make_table([["⛔ Critical error!"]])
            result = table_to_html(table, panel_type="warning")
            assert 'ac:name="warning"' in result
            assert "⛔" not in result
            assert "Critical error!" in result

    def test_note_panel_strips_emoji(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import table_to_html
            table = self._make_table([["⚠️ Note text"]])
            result = table_to_html(table, panel_type="note")
            assert 'ac:name="note"' in result
            assert "Note text" in result

    def test_regular_table_with_cell_color(self):
        with patch("publish_to_confluence.get_cell_color") as mock_gcc:
            from publish_to_confluence import table_to_html
            mock_gcc.return_value = "DCFCE7"  # Green
            table = self._make_table([["H1"], ["D1"]])
            result = table_to_html(table)
            assert "background-color:" in result
            assert "e3fcef" in result

    def test_regular_table_header_gets_default_bg(self):
        with patch("publish_to_confluence.get_cell_color", return_value=None):
            from publish_to_confluence import table_to_html
            table = self._make_table([["Header"], ["Data"]])
            result = table_to_html(table)
            assert "#f4f5f7" in result

    def test_regular_table_data_cell_with_bg(self):
        with patch("publish_to_confluence.get_cell_color") as mock_gcc:
            from publish_to_confluence import table_to_html
            def gcc_side_effect(cell):
                if "Data" in getattr(cell, "text", ""):
                    return "DBEAFE"
                return None
            mock_gcc.side_effect = gcc_side_effect
            table = self._make_table([["H"], ["Data"]])
            result = table_to_html(table)
            assert "deebff" in result


# ── hex_to_confluence_color lowercase input ─────────────────────────────────

class TestHexToConfluenceColorCase:
    def test_lowercase_hex_mapped(self):
        from publish_to_confluence import hex_to_confluence_color
        assert hex_to_confluence_color("ffdd00") == "#fffae6"

    def test_value_error_returns_none(self):
        from publish_to_confluence import hex_to_confluence_color
        assert hex_to_confluence_color("GGGGGG") is None


# ── para_to_html style None ─────────────────────────────────────────────────

class TestParaToHtmlStyleNone:
    def test_para_style_none_uses_normal(self):
        from publish_to_confluence import para_to_html
        para = MagicMock()
        para.text = "Text without style"
        para.style = None
        result = para_to_html(para)
        assert "<p>" in result
        assert "Text without style" in result


# ── main() flow tests ──────────────────────────────────────────────────────

class TestMainFlow:
    """Test main() entry point with mocked dependencies."""

    def test_main_no_token_exits(self):
        from publish_to_confluence import main
        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "", "CONFLUENCE_PERSONAL_TOKEN": ""}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", "x.html", "--project", "P"]):
                with pytest.raises(SystemExit) as exc_info:
                    main()
        assert exc_info.value.code == 1

    def test_main_from_file_no_project_exits(self):
        from publish_to_confluence import main
        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": ""}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", "x.html"]):
                with pytest.raises(SystemExit) as exc_info:
                    main()
        assert exc_info.value.code == 1

    def test_main_docx_mode_no_path_exits(self):
        from publish_to_confluence import main
        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok"}, clear=False):
            with patch("sys.argv", ["publish"]):
                with pytest.raises(SystemExit) as exc_info:
                    main()
        assert exc_info.value.code == 1

    def test_main_docx_mode_docx_not_installed_exits(self, tmp_path):
        docx_path = tmp_path / "fm.docx"
        docx_path.write_bytes(b"dummy")
        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok"}, clear=False):
            with patch("sys.argv", ["publish", str(docx_path)]):
                with patch("publish_to_confluence.docx", None):
                    with pytest.raises(SystemExit) as exc_info:
                        from publish_to_confluence import main
                        main()
        assert exc_info.value.code == 1

    def test_main_from_file_success(self, tmp_path, capsys):
        xhtml_file = tmp_path / "body.xhtml"
        xhtml_file.write_text("<p>Test content</p>")
        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "PRJ"}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", str(xhtml_file), "--project", "PRJ"]):
                with patch("publish_to_confluence._get_page_id", return_value="12345"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>Test</p>", [])):
                            from publish_to_confluence import main
                            main()
        out = capsys.readouterr()
        assert "ГОТОВО" in out.out or "Lock acquired" in out.out

    def test_main_confluence_lock_error_exits(self, tmp_path):
        try:
            from fm_review.confluence_utils import ConfluenceLockError
        except ImportError:
            pytest.skip("fm_review.confluence_utils not available (tenacity?)")
        xhtml_file = tmp_path / "body.xhtml"
        xhtml_file.write_text("<p>Test</p>")
        mock_client = MagicMock()
        mock_client.lock.side_effect = ConfluenceLockError("Locked")

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "PRJ"}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", str(xhtml_file), "--project", "PRJ"]):
                with patch("publish_to_confluence._get_page_id", return_value="12345"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>Test</p>", [])):
                            with pytest.raises(SystemExit) as exc_info:
                                from publish_to_confluence import main
                                main()
        assert exc_info.value.code == 1

    def test_main_confluence_api_error_exits(self, tmp_path):
        try:
            from fm_review.confluence_utils import ConfluenceAPIError
        except ImportError:
            pytest.skip("fm_review.confluence_utils not available (tenacity?)")
        xhtml_file = tmp_path / "body.xhtml"
        xhtml_file.write_text("<p>Test</p>")
        mock_client = MagicMock()
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.side_effect = ConfluenceAPIError("API error", code=500)

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "PRJ"}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", str(xhtml_file), "--project", "PRJ"]):
                with patch("publish_to_confluence._get_page_id", return_value="12345"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>Test</p>", [])):
                            with pytest.raises(SystemExit) as exc_info:
                                from publish_to_confluence import main
                                main()
        assert exc_info.value.code == 1

    def test_main_sanitizer_warnings_printed(self, tmp_path, capsys):
        try:
            import fm_review.confluence_utils  # noqa: F401
        except ImportError:
            pytest.skip("fm_review.confluence_utils not available (tenacity?)")
        xhtml_file = tmp_path / "body.xhtml"
        xhtml_file.write_text("<p>Test</p>")
        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "PRJ"}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", str(xhtml_file), "--project", "PRJ"]):
                with patch("publish_to_confluence._get_page_id", return_value="12345"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>Test</p>", ["Warning 1"])):
                            from publish_to_confluence import main
                            main()
        out = capsys.readouterr()
        assert "Warning 1" in out.out or "Sanitizer" in out.out or "XHTML" in out.out


# ── ConfluenceAPIError with code attribute ─────────────────────────────────

class TestConfluenceAPIErrorWithCode:
    def test_main_api_error_prints_http_code(self, tmp_path, capsys):
        try:
            from fm_review.confluence_utils import ConfluenceAPIError
        except ImportError:
            pytest.skip("fm_review.confluence_utils not available")
        xhtml_file = tmp_path / "body.xhtml"
        xhtml_file.write_text("<p>Test</p>")
        mock_client = MagicMock()
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.side_effect = ConfluenceAPIError("API error", code=500)

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "PRJ"}, clear=False):
            with patch("sys.argv", ["publish", "--from-file", str(xhtml_file), "--project", "PRJ"]):
                with patch("publish_to_confluence._get_page_id", return_value="12345"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>Test</p>", [])):
                            with pytest.raises(SystemExit):
                                from publish_to_confluence import main
                                main()
        out = capsys.readouterr()
        assert "500" in out.out or "HTTP" in out.out


# ── history_table edge case ─────────────────────────────────────────────────

class TestHistoryTableEdgeCase:
    def test_history_table_with_two_rows_not_shortened(self):
        """When len(table.rows) <= 2, history_table_to_html is not used in main loop."""
        from publish_to_confluence import is_history_table, history_table_to_html
        row1 = MagicMock()
        row1.cells = [MagicMock()]
        row1.cells[0].text = "Версия"
        row2 = MagicMock()
        row2.cells = [MagicMock()]
        row2.cells[0].text = "1.0"
        table = MagicMock()
        table.rows = [row1, row2]
        assert is_history_table(table) is True
        result = history_table_to_html(table)
        assert "1.0.0" in result


# ── para_to_html dash prefix ───────────────────────────────────────────────

class TestParaToHtmlDashPrefix:
    def test_dash_prefix_becomes_list_item(self):
        from publish_to_confluence import para_to_html
        para = MagicMock()
        para.text = "- Item with dash"
        para.style = MagicMock()
        para.style.name = "Normal"
        result = para_to_html(para)
        assert "<li>" in result
        assert "Item with dash" in result


# ── main() docx mode (requires python-docx) ───────────────────────────────

class TestMainDocxMode:
    def test_main_docx_mode_full_flow(self, tmp_path, capsys):
        """Test docx import mode with a real docx file."""
        try:
            import docx
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc_path = tmp_path / "FM-TEST-v1.0.0.docx"
        doc = Document()
        doc.add_paragraph("Test FM Title", style="Title")
        doc.add_paragraph("Normal paragraph")
        doc.add_paragraph("Дата последнего изменения: 01.01.2026")
        doc.add_paragraph("- List item")
        doc.save(str(doc_path))

        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        projects_dir = tmp_path / "projects" / "TEST"
        projects_dir.mkdir(parents=True)
        (projects_dir / "CONFLUENCE_PAGE_ID").write_text("99999\n")

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": ""}, clear=False):
            with patch("sys.argv", ["publish", str(doc_path)]):
                with patch("publish_to_confluence._get_page_id", return_value="99999"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml") as mock_sanitize:
                            mock_sanitize.return_value = ("<p>sanitized</p>", [])
                            from publish_to_confluence import main
                            main()
        out = capsys.readouterr()
        assert "ГОТОВО" in out.out or "Lock acquired" in out.out
        assert mock_client.update_page.called

    def test_main_docx_mode_project_from_path(self, tmp_path, capsys):
        """Test docx path with projects/PROJ/ in path triggers project extraction."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        projects_dir = tmp_path / "projects" / "MYPROJ"
        projects_dir.mkdir(parents=True)
        (projects_dir / "CONFLUENCE_PAGE_ID").write_text("77777\n")
        doc_path = projects_dir / "FM-X-v1.0.0.docx"
        doc = Document()
        doc.add_paragraph("Title", style="Title")
        doc.save(str(doc_path))

        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok", "PROJECT": "", "CONFLUENCE_PAGE_ID": ""}, clear=False):
            with patch("sys.argv", ["publish", str(doc_path)]):
                with patch("publish_to_confluence._get_page_id") as mock_get_id:
                    mock_get_id.return_value = "77777"
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml", return_value=("<p>x</p>", [])):
                            from publish_to_confluence import main
                            main()
        mock_get_id.assert_called_once()
        call_args = mock_get_id.call_args[0][0]
        assert call_args == "MYPROJ"

    def test_main_docx_mode_with_code_system_and_tables(self, tmp_path, capsys):
        """Test docx with code system section, meta table, warning paragraph."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc_path = tmp_path / "FM-CODE-v1.0.0.docx"
        doc = Document()
        doc.add_paragraph("FM Title Here", style="Title")
        doc.add_paragraph("Система кодов", style="Heading 2")
        doc.add_paragraph("В документе используется следующая система кодов.")
        doc.add_paragraph("• LS-BR-XXX - Описание кода")
        doc.add_paragraph("Продолжение описания.")
        doc.add_paragraph("Следующий раздел", style="Heading 1")
        doc.add_paragraph("⚠ Важно: проверить настройки")
        doc.add_paragraph("* List item one")
        doc.add_paragraph("- List item two")
        # Meta table: Версия/Дата/Статус/Автор
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "Версия"
        table.rows[0].cells[1].text = "1.0.0"
        table.rows[1].cells[0].text = "Дата"
        table.rows[1].cells[1].text = "01.01.2026"
        table.rows[2].cells[0].text = "Статус"
        table.rows[2].cells[1].text = "DRAFT"
        table.rows[3].cells[0].text = "Автор"
        table.rows[3].cells[1].text = "Test"
        doc.save(str(doc_path))

        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok"}, clear=False):
            with patch("sys.argv", ["publish", str(doc_path)]):
                with patch("publish_to_confluence._get_page_id", return_value="99999"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml") as mock_sanitize:
                            def sanitize_identity(body, *a, **k):
                                return (body, [])
                            mock_sanitize.side_effect = lambda x: (x, [])
                            from publish_to_confluence import main
                            main()
        assert mock_client.update_page.called
        call_body = mock_client.update_page.call_args[1]["new_body"]
        assert "LS-BR-XXX" in call_body or "Код" in call_body
        assert "Важно" in call_body or "note" in call_body
        assert "confluenceTable" in call_body

    def test_main_docx_mode_code_system_table_and_skip(self, tmp_path):
        """Test docx with code system table (Код/Описание) and skip patterns."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc_path = tmp_path / "FM-SKIP-v1.0.0.docx"
        doc = Document()
        doc.add_paragraph("No Title Style", style="Heading 1")  # No Title -> FM_NAME = FM-SKIP
        doc.add_paragraph("Система кодов", style="Heading 2")
        doc.add_paragraph("В документе используется.")
        doc.add_paragraph("• LS-BR-XXX - Code desc.")
        doc.add_paragraph("Other text in section")  # Falls through to continue
        doc.add_paragraph("Next Section", style="Heading 1")
        doc.add_paragraph("* First list")
        doc.add_paragraph("")  # Empty - closes list
        doc.add_paragraph("Маршруты согласования")  # Skip pattern
        # Code system table: Код + Описание
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Код"
        table.rows[0].cells[1].text = "Описание"
        table.rows[1].cells[0].text = "X"
        table.rows[1].cells[1].text = "Y"
        doc.save(str(doc_path))

        mock_client = MagicMock()
        mock_client.get_page.return_value = {"title": "FM", "version": {"number": 1}}
        mock_client.update_page.return_value = ({"version": {"number": 2}}, None)
        mock_lock = MagicMock()
        mock_lock.__enter__ = MagicMock(return_value=None)
        mock_lock.__exit__ = MagicMock(return_value=False)
        mock_client.lock.return_value = mock_lock

        with patch.dict(os.environ, {"CONFLUENCE_TOKEN": "tok"}, clear=False):
            with patch("sys.argv", ["publish", str(doc_path)]):
                with patch("publish_to_confluence._get_page_id", return_value="99999"):
                    with patch("fm_review.confluence_utils.ConfluenceClient", return_value=mock_client):
                        with patch("fm_review.xhtml_sanitizer.sanitize_xhtml") as mock_sanitize:
                            mock_sanitize.side_effect = lambda x: (x, [])
                            from publish_to_confluence import main
                            main()
        assert mock_client.update_page.called


# ── __main__ block ────────────────────────────────────────────────────────

class TestMainBlock:
    def test_main_block_calls_main(self):
        """Verify __name__ == '__main__' block exists and would invoke main()."""
        script_path = SCRIPTS_DIR / "publish_to_confluence.py"
        code = script_path.read_text()
        assert 'if __name__ == "__main__"' in code or "if __name__ == '__main__'" in code
        assert "main()" in code

