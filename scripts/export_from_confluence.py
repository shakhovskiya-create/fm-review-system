#!/usr/bin/env python3
"""
FM Exporter from Confluence v1.0
- Confluence XHTML -> PDF (WeasyPrint)
- Confluence XHTML -> Word (.docx, python-docx)
- Полный контроль шрифтов и форматирования
"""
import json
import os
import re
import sys
import urllib.request
from datetime import datetime

from bs4 import BeautifulSoup
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_exponential

from fm_review.confluence_utils import _get_page_id, _make_ssl_context

# Config - safe module-level defaults (no side effects; validation in main())
CONFLUENCE_URL = os.environ.get("CONFLUENCE_URL", "https://confluence.ekf.su")
TOKEN = os.environ.get("CONFLUENCE_TOKEN", "")


def setup_weasyprint_env():
    """Set DYLD_LIBRARY_PATH on macOS for WeasyPrint and re-exec if needed."""
    if sys.platform == "darwin" and "DYLD_LIBRARY_PATH" not in os.environ:
        homebrew_lib = "/opt/homebrew/lib"
        if os.path.isdir(homebrew_lib):
            os.environ["DYLD_LIBRARY_PATH"] = homebrew_lib
            # Re-exec with the env var set (DYLD_* must be set before process start)
            os.execve(sys.executable, [sys.executable] + sys.argv, os.environ)


# Module-level defaults — set lazily via main() or direct assignment for testing.
# No side effects (no sys.exit, no os.execve) at import time.
PAGE_ID = None

# Output directory (default; can be overridden before calling main)
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "exports")


@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=10),
    retry=retry_if_exception_type(urllib.error.URLError),
    reraise=True,
)
def _urlopen_with_retry(req):
    """urllib.urlopen with tenacity retry on transient errors."""
    return urllib.request.urlopen(req, timeout=30, context=_make_ssl_context())


def api_request(method, endpoint):
    """Confluence REST API request with retry on transient errors."""
    url = f"{CONFLUENCE_URL}/rest/api/{endpoint}"
    req = urllib.request.Request(url, method=method)
    req.add_header("Authorization", f"Bearer {TOKEN}")
    req.add_header("Content-Type", "application/json")
    try:
        with _urlopen_with_retry(req) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except Exception as e:
        print(f"API error: {e}")
        return None


def fetch_page(page_id=None):
    """Fetch page content and metadata from Confluence"""
    pid = page_id or PAGE_ID
    data = api_request("GET", f"content/{pid}?expand=body.storage,version")
    if not data:
        print("Failed to fetch page")
        sys.exit(1)
    return {
        "title": data["title"],
        "version": data["version"]["number"],
        "html": data["body"]["storage"]["value"],
    }


def confluence_to_clean_html(raw_html, title):
    """Convert Confluence XHTML storage format to clean HTML for rendering"""
    soup = BeautifulSoup(raw_html, "html.parser")

    # Expand Confluence macros
    for macro in soup.find_all("ac:structured-macro"):
        macro_name = macro.get("ac:name", "")
        body = macro.find("ac:rich-text-body")
        body_html = body.decode_contents() if body else ""

        if macro_name == "warning":
            # Red panel
            new_tag = soup.new_tag("div")
            new_tag["class"] = "panel-warning"
            new_tag["style"] = (
                "background-color: #ffebe6; border-left: 4px solid #de350b; "
                "padding: 12px 16px; margin: 12px 0; border-radius: 3px;"
            )
            new_tag.append(BeautifulSoup(body_html, "html.parser"))
            macro.replace_with(new_tag)

        elif macro_name == "note":
            # Yellow panel
            new_tag = soup.new_tag("div")
            new_tag["class"] = "panel-note"
            new_tag["style"] = (
                "background-color: #fffae6; border-left: 4px solid #ff991f; "
                "padding: 12px 16px; margin: 12px 0; border-radius: 3px;"
            )
            new_tag.append(BeautifulSoup(body_html, "html.parser"))
            macro.replace_with(new_tag)

        elif macro_name == "info":
            new_tag = soup.new_tag("div")
            new_tag["class"] = "panel-info"
            new_tag["style"] = (
                "background-color: #deebff; border-left: 4px solid #0065ff; "
                "padding: 12px 16px; margin: 12px 0; border-radius: 3px;"
            )
            new_tag.append(BeautifulSoup(body_html, "html.parser"))
            macro.replace_with(new_tag)

        elif macro_name == "toc":
            macro.decompose()
        else:
            # Unknown macro - just keep body
            if body:
                macro.replace_with(BeautifulSoup(body_html, "html.parser"))
            else:
                macro.decompose()

    # Remove ac: namespace elements that aren't macros
    for tag in soup.find_all(re.compile(r"^ac:")):
        tag.unwrap() if tag.string else tag.decompose()

    # Remove ri: namespace elements
    for tag in soup.find_all(re.compile(r"^ri:")):
        tag.decompose()

    content_html = str(soup)

    # Build full HTML document with CSS
    full_html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
    @page {{
        size: A4;
        margin: 20mm 15mm 20mm 15mm;
        @top-center {{
            content: "{title}";
            font-size: 8pt;
            color: #666;
        }}
        @bottom-center {{
            content: "Страница " counter(page) " из " counter(pages);
            font-size: 8pt;
            color: #666;
        }}
    }}
    body {{
        font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
        font-size: 10pt;
        line-height: 1.5;
        color: #172b4d;
        max-width: 100%;
    }}
    h1 {{
        font-size: 18pt;
        font-weight: bold;
        color: #172b4d;
        margin-top: 24px;
        margin-bottom: 12px;
        page-break-after: avoid;
    }}
    h2 {{
        font-size: 14pt;
        font-weight: bold;
        color: #172b4d;
        margin-top: 20px;
        margin-bottom: 10px;
        page-break-after: avoid;
    }}
    h3 {{
        font-size: 12pt;
        font-weight: bold;
        color: #172b4d;
        margin-top: 16px;
        margin-bottom: 8px;
        page-break-after: avoid;
    }}
    p {{
        margin: 6px 0;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin: 12px 0;
        font-size: 9pt;
        page-break-inside: avoid;
    }}
    th, td {{
        border: 1px solid #dfe1e6;
        padding: 6px 8px;
        text-align: left;
        vertical-align: top;
    }}
    th {{
        background-color: #f4f5f7;
        font-weight: bold;
    }}
    tr:nth-child(even) {{
        background-color: #fafbfc;
    }}
    .panel-warning {{
        background-color: #ffebe6;
        border-left: 4px solid #de350b;
        padding: 12px 16px;
        margin: 12px 0;
        border-radius: 3px;
        page-break-inside: avoid;
    }}
    .panel-note {{
        background-color: #fffae6;
        border-left: 4px solid #ff991f;
        padding: 12px 16px;
        margin: 12px 0;
        border-radius: 3px;
        page-break-inside: avoid;
    }}
    .panel-info {{
        background-color: #deebff;
        border-left: 4px solid #0065ff;
        padding: 12px 16px;
        margin: 12px 0;
        border-radius: 3px;
        page-break-inside: avoid;
    }}
    ul, ol {{
        margin: 6px 0 6px 20px;
        padding: 0;
    }}
    li {{
        margin: 3px 0;
    }}
    hr {{
        border: none;
        border-top: 1px solid #dfe1e6;
        margin: 16px 0;
    }}
    strong {{
        font-weight: bold;
    }}
</style>
</head>
<body>
{content_html}
</body>
</html>"""
    return full_html


def export_pdf(html, output_path):
    """Export HTML to PDF using WeasyPrint"""
    try:
        from weasyprint import HTML
        HTML(string=html).write_pdf(output_path)
        return True
    except ImportError:
        print("WeasyPrint not installed. Run: pip3 install weasyprint")
        return False
    except Exception as e:
        print(f"PDF error: {e}")
        return False


def export_docx(raw_html, title, output_path):
    """Export Confluence HTML to Word .docx using python-docx"""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Segoe UI"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x17, 0x2B, 0x4D)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    soup = BeautifulSoup(raw_html, "html.parser")

    def add_text_with_formatting(paragraph, element):
        """Add text from an HTML element preserving bold/italic"""
        if isinstance(element, str):
            text = element.strip()
            if text:
                paragraph.add_run(text)
            return

        if element.name == "strong" or element.name == "b":
            for child in element.children:
                if isinstance(child, str):
                    run = paragraph.add_run(child)
                    run.bold = True
                else:
                    add_text_with_formatting(paragraph, child)
        elif element.name == "em" or element.name == "i":
            for child in element.children:
                if isinstance(child, str):
                    run = paragraph.add_run(child)
                    run.italic = True
                else:
                    add_text_with_formatting(paragraph, child)
        elif element.name == "br":
            paragraph.add_run("\n")
        elif element.name in ("span", "a", "u"):
            for child in element.children:
                add_text_with_formatting(paragraph, child)
        else:
            for child in element.children:
                add_text_with_formatting(paragraph, child)

    def process_element(element):
        """Process an HTML element and add to Word document"""
        if isinstance(element, str):
            text = element.strip()
            if text:
                p = doc.add_paragraph()
                p.add_run(text)
            return

        tag = element.name
        if not tag:
            return

        # Headings
        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            heading = doc.add_heading(level=level)
            add_text_with_formatting(heading, element)
            return

        # Paragraphs
        if tag == "p":
            text = element.get_text(strip=True)
            if not text:
                return
            p = doc.add_paragraph()
            add_text_with_formatting(p, element)
            return

        # Tables
        if tag == "table":
            rows = element.find_all("tr")
            if not rows:
                return

            # Count max columns
            max_cols = 0
            for row in rows:
                cells = row.find_all(["th", "td"])
                max_cols = max(max_cols, len(cells))

            if max_cols == 0:
                return

            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for j, cell in enumerate(cells):
                    if j >= max_cols:
                        break
                    doc_cell = table.rows[i].cells[j]
                    doc_cell.text = ""
                    p = doc_cell.paragraphs[0]
                    add_text_with_formatting(p, cell)
                    p.paragraph_format.space_after = Pt(2)

                    # Header row styling
                    if cell.name == "th" or i == 0:
                        for run in p.runs:
                            run.bold = True

                    # Background color from style
                    style_attr = cell.get("style", "")
                    bg_match = re.search(
                        r"background-color:\s*([#\w]+)", style_attr
                    )
                    if bg_match:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn

                        color = bg_match.group(1).lstrip("#")
                        if len(color) == 6:
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:fill"), color.upper())
                            shading.set(qn("w:val"), "clear")
                            doc_cell._tc.get_or_add_tcPr().append(shading)

            doc.add_paragraph()  # spacing after table
            return

        # Panels (warning, note, info) - converted divs
        if tag == "div":
            cls = element.get("class", [])
            if isinstance(cls, str):
                cls = [cls]

            panel_type = None
            if "panel-warning" in cls:
                panel_type = "warning"
            elif "panel-note" in cls:
                panel_type = "note"
            elif "panel-info" in cls:
                panel_type = "info"

            if panel_type:
                # Create a single-cell table for the panel
                table = doc.add_table(rows=1, cols=1)
                table.style = "Table Grid"
                cell = table.rows[0].cells[0]
                cell.text = ""
                p = cell.paragraphs[0]

                text = element.get_text(strip=True)
                run = p.add_run(text)
                run.font.size = Pt(10)

                # Set background color
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                colors = {
                    "warning": "FFEBE6",
                    "note": "FFFAE6",
                    "info": "DEEBFF",
                }
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), colors.get(panel_type, "F4F5F7"))
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)

                doc.add_paragraph()
                return

            # Regular div - process children
            for child in element.children:
                process_element(child)
            return

        # Lists
        if tag in ("ul", "ol"):
            items = element.find_all("li", recursive=False)
            for item in items:
                p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                add_text_with_formatting(p, item)
            return

        # HR
        if tag == "hr":
            p = doc.add_paragraph()
            p.add_run("_" * 80)
            p.runs[0].font.color.rgb = RGBColor(0xDF, 0xE1, 0xE6)
            return

        # Other block elements - process children
        if tag in ("div", "section", "article", "main", "body", "tbody", "thead"):
            for child in element.children:
                process_element(child)
            return

    # Process all top-level elements
    for element in soup.children:
        process_element(element)

    doc.save(output_path)
    return True


def main():
    global PAGE_ID, TOKEN, CONFLUENCE_URL

    # --- Side effects: only run when executed as a script ---
    # SSL is now per-request via _make_ssl_context() — no global override needed
    setup_weasyprint_env()

    # Re-read env vars (may have changed after re-exec)
    CONFLUENCE_URL = os.environ.get("CONFLUENCE_URL", "https://confluence.ekf.su")
    TOKEN = os.environ.get("CONFLUENCE_TOKEN", "")

    # Handle --help before any validation (no side effects)
    if "--help" in sys.argv or "-h" in sys.argv:
        print("Usage: export_from_confluence.py [--pdf|--docx|--both] [--page=ID] [--project=PROJECT_NAME]")
        print("  --pdf      Export PDF only")
        print("  --docx     Export Word only")
        print("  --both     Export both (default)")
        print("  --page=ID  Page ID (overrides project file)")
        print("  --project  Read PAGE_ID from projects/PROJECT_NAME/CONFLUENCE_PAGE_ID")
        print("  Env: CONFLUENCE_PAGE_ID, CONFLUENCE_TOKEN, CONFLUENCE_URL")
        sys.exit(0)

    if not TOKEN:
        print("ERROR: CONFLUENCE_TOKEN environment variable not set")
        sys.exit(1)

    PAGE_ID = _get_page_id(os.environ.get("PROJECT"))

    # Parse arguments
    fmt = "both"  # pdf, docx, both
    page_id = PAGE_ID
    project_arg = None

    for arg in sys.argv[1:]:
        if arg in ("--pdf", "-p"):
            fmt = "pdf"
        elif arg in ("--docx", "-w", "--word"):
            fmt = "docx"
        elif arg in ("--both", "-b"):
            fmt = "both"
        elif arg.startswith("--page="):
            page_id = arg.split("=")[1]
        elif arg.startswith("--project="):
            project_arg = arg.split("=")[1]
        elif arg in ("--help", "-h"):
            pass  # handled before TOKEN/PAGE_ID checks at start of main()

    if project_arg:
        page_id = _get_page_id(project_arg)

    print("=" * 60)
    print("FM EXPORTER - CONFLUENCE v1.0")
    print("=" * 60)

    # Create output dir
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Fetch page
    print("\n=== ЗАГРУЗКА СТРАНИЦЫ ===")
    page = fetch_page(page_id)
    print(f"  Название: {page['title']}")
    print(f"  Версия: {page['version']}")
    print(f"  HTML: {len(page['html'])} символов")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w\-]", "_", page["title"])

    # Export PDF
    if fmt in ("pdf", "both"):
        print("\n=== ЭКСПОРТ PDF ===")
        clean_html = confluence_to_clean_html(page["html"], page["title"])
        pdf_path = os.path.join(OUTPUT_DIR, f"{safe_title}_v{page['version']}_{timestamp}.pdf")
        if export_pdf(clean_html, pdf_path):
            size_kb = os.path.getsize(pdf_path) / 1024
            print(f"  PDF: {pdf_path}")
            print(f"  Размер: {size_kb:.0f} KB")
        else:
            print("  PDF: ОШИБКА")

    # Export Word
    if fmt in ("docx", "both"):
        print("\n=== ЭКСПОРТ WORD ===")
        docx_path = os.path.join(OUTPUT_DIR, f"{safe_title}_v{page['version']}_{timestamp}.docx")
        # For Word we need to process the raw HTML differently (with panels converted)
        clean_for_word = confluence_to_clean_html(page["html"], page["title"])
        soup = BeautifulSoup(clean_for_word, "html.parser")
        body = soup.find("body")
        body_html = str(body) if body else clean_for_word

        if export_docx(body_html, page["title"], docx_path):
            size_kb = os.path.getsize(docx_path) / 1024
            print(f"  Word: {docx_path}")
            print(f"  Размер: {size_kb:.0f} KB")
        else:
            print("  Word: ОШИБКА")

    print(f"\nГОТОВО! Файлы в: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
