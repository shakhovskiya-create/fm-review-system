#!/usr/bin/env python3
"""
FM Publisher v4
- Требования → БД с полным описанием (для ТЗ)
- Глоссарий, Риски, Версии → таблицы в теле
- Таблицы 1x1 → callout блоки (КРИТИЧЕСКАЯ ЗАВИСИМОСТЬ)
"""
import json, urllib.request, ssl, time, re, sys
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ssl._create_default_https_context = ssl._create_unverified_context

import os
CONFIG_PATH = os.path.join(os.path.dirname(__file__), "../docs/notion-config.json")
with open(CONFIG_PATH) as f:
    CFG = json.load(f)

TOKEN = CFG.get("token", os.environ.get("NOTION_TOKEN", ""))
FM_DB = CFG["fm_database"]
REQ_DB = CFG["requirements_database"]
VER_DB = CFG["versions_database"]

if not TOKEN:
    print("ОШИБКА: Не найден Notion API token в config или переменной окружения NOTION_TOKEN")
    sys.exit(1)

HEADERS = {"Authorization": f"Bearer {TOKEN}", "Content-Type": "application/json", "Notion-Version": "2022-06-28"}

# Get document path from command line or use default
if len(sys.argv) > 1:
    DOC_PATH = sys.argv[1]
else:
    DOC_PATH = "/Users/antonsahovskii/Documents/claude-agents/fm-review-system/projects/PROJECT_SHPMNT_PROFIT/FM_DOCUMENTS/FM-LS-PROFIT-v1.2.1.docx"

def extract_fm_metadata(doc_path):
    """Extract FM code, name, version from filename and document"""
    import os
    filename = os.path.basename(doc_path)
    # Parse filename: FM-CODE-vX.Y.Z.docx
    match = re.match(r'(FM-[A-Z-]+)-v(\d+\.\d+\.\d+)\.docx', filename)
    if match:
        code = match.group(1)
        version = match.group(2)
    else:
        code = "FM-UNKNOWN"
        version = "1.0.0"
    return code, version

def api(method, path, data=None):
    url = f"https://api.notion.com/v1/{path}"
    body = json.dumps(data).encode() if data else None
    req = urllib.request.Request(url, data=body, headers=HEADERS, method=method)
    for attempt in range(3):
        try:
            with urllib.request.urlopen(req, timeout=90) as resp:
                return json.loads(resp.read())
        except Exception as e:
            if attempt < 2:
                time.sleep(2)
            else:
                print(f"  API ERROR: {str(e)[:50]}", file=sys.stderr)
    return None

def rich_text(text, max_len=2000, bold=False, color=None):
    """
    Colors: gray, brown, orange, yellow, green, blue, purple, pink, red
    Background: gray_background, brown_background, etc.
    """
    if not text:
        return []
    rt = {"text": {"content": str(text)[:max_len]}}
    annotations = {}
    if bold:
        annotations["bold"] = True
    if color:
        annotations["color"] = color
    if annotations:
        rt["annotations"] = annotations
    return [rt]

def query_db(db_id):
    pages = []
    r = api("POST", f"databases/{db_id}/query", {"page_size": 100})
    if r:
        pages = r.get("results", [])
    return pages

def archive_page(page_id):
    return api("PATCH", f"pages/{page_id}", {"archived": True})

# === LOAD DOCUMENT ===
print("=" * 60)
print("FM PUBLISHER v5")
print("  - Требования → БД")
print("  - Глоссарий, Риски → таблицы в теле")
print("  - Цвета из Word → цвет текста")
print("=" * 60)
doc = docx.Document(DOC_PATH)
FM_CODE, FM_VERSION = extract_fm_metadata(DOC_PATH)

# Extract title from document (first non-empty Title paragraph)
FM_NAME = ""
for para in doc.paragraphs:
    if para.style and para.style.name == "Title" and para.text.strip():
        FM_NAME = para.text.strip()
        break
if not FM_NAME:
    FM_NAME = FM_CODE

print(f"Документ: {DOC_PATH.split('/')[-1]}")
print(f"Код: {FM_CODE}, Версия: {FM_VERSION}")
print(f"Название: {FM_NAME[:50]}...")

# === STEP 0: FIND EXISTING FM OR PREPARE TO CREATE ===
print("\n=== ПОИСК СУЩЕСТВУЮЩЕЙ ФМ ===")

def find_fm_page(code):
    """Find existing FM page by code"""
    pages = query_db(FM_DB)
    for p in pages:
        props = p.get("properties", {})
        code_prop = props.get("Код", {}).get("title", [{}])
        code_text = code_prop[0].get("text", {}).get("content", "") if code_prop else ""
        if code_text == code:
            return p
    return None

def clear_page_content(page_id):
    """Remove all blocks from page"""
    r = api("GET", f"blocks/{page_id}/children?page_size=100")
    if r:
        for block in r.get("results", []):
            api("DELETE", f"blocks/{block['id']}")
            time.sleep(0.05)

def archive_related_records(fm_id, db_id):
    """Archive records linked to FM"""
    pages = query_db(db_id)
    count = 0
    for p in pages:
        props = p.get("properties", {})
        fm_rel = props.get("ФМ", {}).get("relation", [])
        for rel in fm_rel:
            if rel.get("id") == fm_id:
                archive_page(p["id"])
                count += 1
                time.sleep(0.05)
                break
    return count

existing_fm = find_fm_page(FM_CODE)
if existing_fm:
    print(f"  Найдена существующая ФМ: {FM_CODE}")
    FM_ID = existing_fm["id"]
    FM_URL = existing_fm["url"]

    # Clear old content
    print("  Очищаю контент страницы...")
    clear_page_content(FM_ID)

    # Archive old requirements and versions for this FM
    print("  Архивирую старые требования и версии...")
    archived = archive_related_records(FM_ID, REQ_DB)
    archived += archive_related_records(FM_ID, VER_DB)
    print(f"  Архивировано: {archived}")

    # Update FM properties
    api("PATCH", f"pages/{FM_ID}", {
        "properties": {
            "Название": {"rich_text": rich_text(FM_NAME)},
            "Версия": {"rich_text": rich_text(FM_VERSION)},
        }
    })
    print(f"  Обновлена версия: {FM_VERSION}")
else:
    print(f"  ФМ не найдена, будет создана новая")

# === STEP 1: CREATE FM PAGE (if not exists) ===
if not existing_fm:
    print("\n=== СОЗДАНИЕ FM ===")
    fm_page = api("POST", "pages", {
        "parent": {"database_id": FM_DB},
        "properties": {
            "Код": {"title": rich_text(FM_CODE)},
            "Название": {"rich_text": rich_text(FM_NAME)},
            "Версия": {"rich_text": rich_text(FM_VERSION)},
            "Статус": {"select": {"name": "Draft"}},
            "Приоритет": {"select": {"name": "P1"}},
            "Область": {"multi_select": [{"name": "Продажи"}, {"name": "Логистика"}, {"name": "Финансы"}]},
            "Системы": {"multi_select": [{"name": "1С:УТ"}, {"name": "1С:ERP"}]}
        }
    })
    if not fm_page:
        print("ОШИБКА!")
        sys.exit(1)
    FM_ID = fm_page["id"]
    FM_URL = fm_page["url"]
    print(f"  ID: {FM_ID}")
    print(f"  URL: {FM_URL}")
else:
    print(f"\n=== ОБНОВЛЕНИЕ FM ===")
    print(f"  ID: {FM_ID}")
    print(f"  URL: {FM_URL}")

# === STEP 2: CREATE REQUIREMENTS IN DB ===
print("\n=== СОЗДАНИЕ ТРЕБОВАНИЙ В БД ===")

def find_requirements_table(doc):
    """Find requirements table by pattern (first cell contains XX-YY-NNN)"""
    for idx, table in enumerate(doc.tables):
        if len(table.rows) > 1:
            first_cell = table.rows[1].cells[0].text.strip() if table.rows[1].cells else ""
            # Check for requirement code pattern
            if re.match(r'[A-Z]+-[A-Z]+-\d+', first_cell):
                print(f"  Найдена таблица требований: индекс {idx}")
                return table
    return None

req_table = find_requirements_table(doc)
if not req_table:
    print("  ВНИМАНИЕ: Таблица требований не найдена!")
    req_count = 0
else:
    req_count = 0
    for row in req_table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 3 and cells[0]:
            code = cells[0]
            name = cells[1]
            description = cells[2]  # Full description
            type_match = re.search(r'LS-(BR|FR|WF|RPT|NFR|INT|SEC)', code)
            req_type = type_match.group(1) if type_match else "FR"

            # Create requirement with description in DB field
            api("POST", "pages", {
                "parent": {"database_id": REQ_DB},
                "properties": {
                    "Код": {"title": rich_text(code)},
                    "Название": {"rich_text": rich_text(name)},
                    "Описание": {"rich_text": rich_text(description)},
                    "Тип": {"select": {"name": req_type}},
                    "Статус": {"select": {"name": "New"}},
                    "Приоритет": {"select": {"name": "P1 (MVP)"}},
                    "ФМ": {"relation": [{"id": FM_ID}]}
                }
            })
            req_count += 1
            time.sleep(0.15)
print(f"  Создано: {req_count} требований с описанием")

# === STEP 2b: CREATE INITIAL VERSION IN DB ===
print("\n=== СОЗДАНИЕ НАЧАЛЬНОЙ ВЕРСИИ В БД ===")
import datetime
today = datetime.date.today().isoformat()

api("POST", "pages", {
    "parent": {"database_id": VER_DB},
    "properties": {
        "Версия": {"title": rich_text(f"v{FM_VERSION}")},
        "Дата": {"date": {"start": today}},
        "Изменения": {"rich_text": rich_text("Начальная версия документа в Notion")},
        "Тип": {"select": {"name": "Major"}},
        "ФМ": {"relation": [{"id": FM_ID}]}
    }
})
print(f"  Создана версия: v{FM_VERSION} (начальная)")

# === STEP 3: BUILD CONTENT ===
print("\n=== ПОСТРОЕНИЕ КОНТЕНТА ===")

# Tables that become linked databases (placeholder only)
LINKED_DB_TABLES = {
    1: "История версий",   # Table 1 = version history
}
# Note: Table 24 (requirements) is shown as regular table for readability
# Requirements are ALSO created in DB for ТЗ generation

# Tables to skip (redundant with page properties)
SKIP_TABLES = {0}  # Table 0 = document passport

def hex_to_notion_color(hex_color):
    """Map Word hex color to Notion text color"""
    if not hex_color or hex_color == 'auto' or hex_color == 'none':
        return None
    hex_color = hex_color.upper()

    # Exact matches for common Tailwind-like colors
    COLOR_MAP = {
        'FFDD00': 'yellow',   # Header yellow
        'DCFCE7': 'green',    # Green-100 (light green)
        'FEF3C7': 'yellow',   # Amber-100 (light yellow)
        'FED7AA': 'orange',   # Orange-200
        'FECACA': 'red',      # Red-200 (light red)
        'DBEAFE': 'blue',     # Blue-100 (light blue)
        'FEE2E2': 'red',      # Red-100
        'D1FAE5': 'green',    # Green-200
        'FDE68A': 'yellow',   # Yellow-300
        'FBBF24': 'yellow',   # Yellow-400
        'F87171': 'red',      # Red-400
        '10B981': 'green',    # Green-500
    }

    if hex_color in COLOR_MAP:
        return COLOR_MAP[hex_color]

    # Fallback: analyze hex components
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)

    # Green-ish (g > r and g > b)
    if g > r and g > b and g > 180:
        return "green"
    # Red-ish (r > g and r > b, but not yellow)
    if r > g and r > b and r > 200 and g < 200:
        return "red"
    # Yellow-ish (r and g both high, b low)
    if r > 200 and g > 180 and b < 150:
        return "yellow"
    # Blue-ish
    if b > r and b > g and b > 180:
        return "blue"
    # Orange (r high, g medium, b low)
    if r > 220 and 100 < g < 220 and b < 180:
        return "orange"

    return None

def get_cell_color(cell):
    """Extract fill color from Word cell"""
    try:
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                return shd.get(qn('w:fill'))
    except:
        pass
    return None

def make_heading(level, text):
    key = f"heading_{level}"
    return {"object": "block", "type": key, key: {"rich_text": rich_text(text)}}

def make_paragraph(text):
    if not text or not text.strip():
        return None
    return {"object": "block", "type": "paragraph", "paragraph": {"rich_text": rich_text(text)}}

def make_callout(text, emoji, color):
    clean = text.lstrip("⚠️⛔❌✅ ")
    return {"object": "block", "type": "callout", "callout": {
        "rich_text": rich_text(clean),
        "icon": {"type": "emoji", "emoji": emoji},
        "color": color
    }}

def make_bullet(text):
    return {"object": "block", "type": "bulleted_list_item", "bulleted_list_item": {"rich_text": rich_text(text.lstrip("•- "))}}

def make_divider():
    return {"object": "block", "type": "divider", "divider": {}}

def make_table_block(table):
    # Collect data with colors
    rows_data = []
    rows_colors = []
    for row in table.rows:
        row_data = []
        row_color = []
        for cell in row.cells:
            row_data.append(cell.text.strip()[:500])
            row_color.append(get_cell_color(cell))
        if any(cell for cell in row_data):
            rows_data.append(row_data)
            rows_colors.append(row_color)

    if len(rows_data) < 1:
        return None

    # Check if it's a 1x1 table (callout-style)
    if len(rows_data) == 1 and len(rows_data[0]) == 1:
        text = rows_data[0][0]
        if "⛔" in text or "КРИТИЧЕСК" in text.upper():
            return make_callout(text, "⛔", "red_background")
        elif "⚠" in text:
            return make_callout(text, "⚠️", "yellow_background")

    if len(rows_data) < 2:
        # Single row - just paragraph
        return make_paragraph(" | ".join(rows_data[0]))

    width = max(len(row) for row in rows_data)
    table_rows = []
    for row_idx, row_data in enumerate(rows_data):
        cells = []
        for i in range(width):
            cell_text = row_data[i] if i < len(row_data) else ""
            cell_hex = rows_colors[row_idx][i] if i < len(rows_colors[row_idx]) else None
            notion_color = hex_to_notion_color(cell_hex)
            # If has color - apply bold + color
            if notion_color:
                cells.append(rich_text(str(cell_text), bold=True, color=notion_color))
            else:
                cells.append(rich_text(str(cell_text)))
        table_rows.append({
            "object": "block",
            "type": "table_row",
            "table_row": {"cells": cells}
        })

    return {
        "object": "block",
        "type": "table",
        "table": {
            "table_width": width,
            "has_column_header": True,
            "has_row_header": False,
            "children": table_rows
        }
    }

def make_linked_db_placeholder(db_name):
    """Placeholder for linked database - user will replace with /linked"""
    emoji = "📋" if db_name == "Требования" else "📜"
    return {
        "object": "block",
        "type": "callout",
        "callout": {
            "rich_text": rich_text(f"👆 Замени на /linked → {db_name} → скрой колонку ФМ"),
            "icon": {"type": "emoji", "emoji": emoji},
            "color": "blue_background"
        }
    }

# Parse document in order
body = doc.element.body
blocks = []
table_idx = 0

for child in body:
    if isinstance(child, CT_P):
        para = Paragraph(child, doc)
        text = para.text.strip()
        if not text:
            continue

        # Skip redundant metadata (already in page properties)
        if text.startswith("Дата последнего изменения"):
            continue

        style = para.style.name if para.style else "Normal"

        if text.startswith("⛔"):
            blocks.append(make_callout(text, "⛔", "red_background"))
        elif text.startswith("⚠"):
            blocks.append(make_callout(text, "⚠️", "yellow_background"))
        elif style == "Title":
            if "КОНТРОЛЬ РЕНТАБЕЛЬНОСТИ" in text:
                blocks.append(make_heading(1, text))
        elif style == "Heading 1":
            blocks.append(make_heading(1, text))
        elif style == "Heading 2":
            blocks.append(make_heading(2, text))
        elif style == "Heading 3":
            blocks.append(make_heading(3, text))
        elif style == "List Paragraph" or text.startswith("•") or text.startswith("-"):
            blocks.append(make_bullet(text))
        else:
            b = make_paragraph(text)
            if b:
                blocks.append(b)

    elif isinstance(child, CT_Tbl):
        table = Table(child, doc)

        if table_idx in SKIP_TABLES:
            # Skip passport table (redundant with page properties)
            pass
        elif table_idx in LINKED_DB_TABLES:
            # Linked database placeholder
            db_name = LINKED_DB_TABLES[table_idx]
            blocks.append(make_linked_db_placeholder(db_name))
        else:
            # Regular table (glossary, risks, etc.) with colors
            tb = make_table_block(table)
            if tb:
                blocks.append(tb)

        table_idx += 1

# Filter None
blocks = [b for b in blocks if b]
print(f"  Блоков: {len(blocks)}")

# === STEP 4: ADD BLOCKS ===
print("\n=== ДОБАВЛЕНИЕ КОНТЕНТА ===")
batch_size = 30
for i in range(0, len(blocks), batch_size):
    chunk = blocks[i:i+batch_size]
    r = api("PATCH", f"blocks/{FM_ID}/children", {"children": chunk})
    if r:
        print(f"  Batch {i//batch_size + 1}: +{len(chunk)}")
    else:
        print(f"  Batch {i//batch_size + 1}: ОШИБКА")
    time.sleep(1)

# === DONE ===
print("\n" + "=" * 60)
print("ГОТОВО!")
print("=" * 60)
print(f"FM страница: {FM_URL}")
print(f"Требования:  {req_count} (в БД с описанием)")
print(f"Глоссарий:   таблица в теле")
print(f"Риски:       таблица в теле")
print("=" * 60)
print("\nОсталось:")
print("1. Открой страницу")
print("2. Найди синий блок 'Замени на /linked'")
print("3. Удали его и введи /linked → Требования")
print("4. Скрой колонку ФМ в view")
